# -*- coding: utf-8 -*-
"""
utils/report_export.py

Convertit un rapport texte Markdown (généré par le LLM : #/##/### titres,
**gras**, listes "- ...", tableaux "| a | b |") en document Word (.docx)
ou PDF, prêts à être téléchargés depuis l'app.

Les tableaux Markdown sont détectés et rendus comme de VRAIS tableaux
(colonnes, bordures, en-tête en gras) plutôt que des lignes de texte à plat.
"""

import io
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from fpdf import FPDF
from fpdf.fonts import FontFace

OOREDOO_RED = RGBColor(0xE4, 0x03, 0x2E)
OOREDOO_RED_RGB = (228, 3, 46)


# ---------------------------------------------------------------------------
# ANALYSE DU MARKDOWN EN BLOCS (texte / titre / liste / tableau)
# ---------------------------------------------------------------------------
def _is_markdown_header(line: str) -> bool:
    return bool(re.match(r"^#{1,6}\s+", line)) or (line.startswith("**") and line.endswith("**") and len(line) > 4)


def _strip_header_markers(line: str) -> str:
    line = re.sub(r"^#{1,6}\s+", "", line)
    return line.strip("*").strip()


def _is_bullet(line: str) -> bool:
    return line.startswith("- ") or line.startswith("* ")


def _is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def _is_table_separator(line: str) -> bool:
    """Ligne du type '|---|---|' ou '| :--- | ---: |' qui sépare l'en-tête du corps."""
    if not _is_table_row(line):
        return False
    cells = [c.strip() for c in line.strip("|").split("|")]
    return all(re.fullmatch(r":?-+:?", c) for c in cells if c != "")


def _is_divider(line: str) -> bool:
    """Ligne horizontale Markdown ("---")."""
    return line.strip("-") == "" and line != ""


def _strip_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    return text.strip()


def _table_row_cells(line: str) -> list[str]:
    return [_strip_inline_markdown(c.strip()) for c in line.strip("|").split("|")]


_SELF_REFERENCE_PATTERNS = re.compile(
    r"d[ée]j[àa]\s+list|voir\s+ligne|voir\s+plus\s+haut|voir\s+ci-?dessus|"
    r"already\s+listed|see\s+above|cf\.?\s+ligne",
    re.IGNORECASE,
)


def _clean_table_rows(rows: list[list[str]]) -> list[list[str]]:
    """
    Filet de sécurité déterministe (indépendant du LLM) : supprime les lignes
    fantômes qu'un LLM ajoute parfois malgré les instructions du prompt
    - lignes qui se référencent elles-mêmes ("déjà listé ci-dessus", etc.)
    - doublons de la même entité (même 1ère colonne) apparue plus haut dans
      le même tableau
    """
    seen_keys = set()
    cleaned = []
    for row in rows:
        if not row or not row[0].strip():
            continue
        first_cell = row[0]

        if _SELF_REFERENCE_PATTERNS.search(first_cell):
            continue

        # clé de dédoublonnage : 1ère colonne, sans astérisque/parenthèses/espaces multiples
        key = re.sub(r"\(.*?\)", "", first_cell)
        key = re.sub(r"[\*\s]+", " ", key).strip().lower()
        if key in seen_keys:
            continue
        seen_keys.add(key)
        cleaned.append(row)
    return cleaned


def parse_report_blocks(report_text: str) -> list[dict]:
    """
    Découpe le rapport en une liste de blocs typés :
      {"type": "header", "text": "..."}
      {"type": "bullet", "text": "..."}
      {"type": "text", "text": "..."}
      {"type": "table", "header": [...], "rows": [[...], [...]]}
    """
    lines = [line.strip() for line in report_text.split("\n") if line.strip()]
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        if _is_divider(line) and not _is_table_row(line):
            i += 1
            continue

        # Détection d'un tableau : ligne "| ... |" suivie d'une ligne séparatrice "|---|---|"
        if _is_table_row(line) and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header = _table_row_cells(line)
            rows = []
            i += 2
            while i < len(lines) and _is_table_row(lines[i]):
                rows.append(_table_row_cells(lines[i]))
                i += 1
            blocks.append({"type": "table", "header": header, "rows": _clean_table_rows(rows)})
            continue

        if _is_markdown_header(line):
            blocks.append({"type": "header", "text": _strip_header_markers(line)})
        elif _is_bullet(line):
            blocks.append({"type": "bullet", "text": _strip_inline_markdown(line[2:])})
        else:
            blocks.append({"type": "text", "text": _strip_inline_markdown(line)})
        i += 1

    return blocks


# ---------------------------------------------------------------------------
# EXPORT WORD (.docx)
# ---------------------------------------------------------------------------
def _set_docx_cell_shading(cell, color_hex: str):
    shd = cell._tc.get_or_add_tcPr()
    shd_elem = shd.makeelement(qn("w:shd"), {qn("w:fill"): color_hex})
    shd.append(shd_elem)


def generate_docx_report(report_text: str, title: str) -> bytes:
    doc = Document()

    heading = doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.color.rgb = OOREDOO_RED

    subtitle = doc.add_paragraph("Généré par l'Assistant IA — Ooredoo Sales Intelligence")
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    for block in parse_report_blocks(report_text):
        if block["type"] == "header":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = OOREDOO_RED
            p.space_before = Pt(10)

        elif block["type"] == "bullet":
            doc.add_paragraph(block["text"], style="List Bullet")

        elif block["type"] == "text":
            doc.add_paragraph(block["text"])

        elif block["type"] == "table":
            n_cols = len(block["header"])
            table = doc.add_table(rows=1, cols=n_cols)
            table.style = "Table Grid"

            for j, cell_text in enumerate(block["header"]):
                cell = table.rows[0].cells[j]
                cell.text = cell_text
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_docx_cell_shading(cell, "E4032E")

            for row_data in block["rows"]:
                row_cells = table.add_row().cells
                for j, cell_text in enumerate(row_data[:n_cols]):
                    row_cells[j].text = cell_text
                    for run in row_cells[j].paragraphs[0].runs:
                        run.font.size = Pt(10)

            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# EXPORT PDF
# ---------------------------------------------------------------------------
def _latin1(text: str) -> str:
    """Les polices de base de fpdf2 (Helvetica) sont en Latin-1 : les accents
    français passent très bien, on neutralise juste les caractères hors
    Latin-1 (emojis, tirets typographiques Unicode, espaces insécables, etc.)."""
    return text.encode("latin-1", "replace").decode("latin-1")


def generate_pdf_report(report_text: str, title: str) -> bytes:
    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    pdf.set_fill_color(*OOREDOO_RED_RGB)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 16, _latin1(title), new_x="LMARGIN", new_y="NEXT", fill=True)
    pdf.ln(4)

    pdf.set_text_color(90, 90, 90)
    pdf.set_font("Helvetica", "I", 9)
    pdf.cell(0, 6, "Genere par l'Assistant IA - Ooredoo Sales Intelligence", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_text_color(20, 20, 20)

    for block in parse_report_blocks(report_text):
        if block["type"] == "header":
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(*OOREDOO_RED_RGB)
            pdf.ln(2)
            pdf.multi_cell(0, 7, _latin1(block["text"]), new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(20, 20, 20)

        elif block["type"] == "bullet":
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, "  -  " + _latin1(block["text"]), new_x="LMARGIN", new_y="NEXT")

        elif block["type"] == "text":
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, _latin1(block["text"]), new_x="LMARGIN", new_y="NEXT")

        elif block["type"] == "table":
            n_cols = len(block["header"])
            font_size = 9 if n_cols <= 5 else 7  # tableaux larges -> police plus petite pour tenir en page
            pdf.set_font("Helvetica", "", font_size)
            header_style = FontFace(emphasis="B", color=(255, 255, 255), fill_color=OOREDOO_RED_RGB)
            with pdf.table(
                borders_layout="ALL",
                text_align="LEFT",
                line_height=6,
                col_widths=tuple([1] * n_cols),
                headings_style=header_style,
            ) as pdf_table:
                header_row = pdf_table.row()
                for cell_text in block["header"]:
                    header_row.cell(_latin1(cell_text))
                for row_data in block["rows"]:
                    data_row = pdf_table.row()
                    for cell_text in row_data[:n_cols]:
                        data_row.cell(_latin1(cell_text))
            pdf.ln(3)

    output = pdf.output()
    return bytes(output)